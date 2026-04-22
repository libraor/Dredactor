"""Word文档解析器 - 解析Word文档内容并保留格式信息"""

import os
from typing import List, Optional

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table as DocxTable, _Cell
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from .models import ParsedDocument, TextBlock, Table, TableCell


class DocumentParser:
    """Word文档解析器

    功能：
    - 读取Word文档(.docx)
    - 提取段落文本及格式
    - 提取表格内容
    - 提取页眉页脚
    - 保留原始格式信息
    """

    def __init__(self, extract_comments: bool = False):
        """
        初始化解析器

        Args:
            extract_comments: 是否提取批注（默认False）
        """
        self.extract_comments = extract_comments

    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析Word文档

        Args:
            file_path: Word文档路径

        Returns:
            ParsedDocument: 解析后的文档对象

        Raises:
            FileNotFoundError: 文件不存在
            Exception: 解析失败
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if not file_path.lower().endswith(".docx"):
            raise ValueError("仅支持.docx格式的Word文档")

        try:
            # 加载Word文档
            docx = Document(file_path)

            # 创建ParsedDocument对象
            parsed_doc = ParsedDocument(file_path=file_path, title=self._get_document_title(docx))

            # 提取页眉页脚
            parsed_doc.headers = self._extract_headers(docx)
            parsed_doc.footers = self._extract_footers(docx)

            # 提取段落和表格
            self._extract_content(docx, parsed_doc)

            # 提取文档元数据
            parsed_doc.metadata = self._extract_metadata(docx)

            return parsed_doc

        except Exception as e:
            raise Exception(f"文档解析失败: {str(e)}")

    def _get_document_title(self, docx: DocxDocument) -> str:
        """
        获取文档标题

        优先级：文档属性标题 > 第一个非空段落
        """
        # 尝试从文档属性获取
        if docx.core_properties.title:
            return docx.core_properties.title

        # 使用第一个非空段落作为标题
        for para in docx.paragraphs:
            if para.text.strip():
                return para.text.strip()

        return "未命名文档"

    def _extract_headers(self, docx: DocxDocument) -> List[str]:
        """提取页眉"""
        headers = []

        for section in docx.sections:
            if section.header and section.header.paragraphs:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        headers.append(text)

        return headers

    def _extract_footers(self, docx: DocxDocument) -> List[str]:
        """提取页脚"""
        footers = []

        for section in docx.sections:
            if section.footer and section.footer.paragraphs:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        footers.append(text)

        return footers

    def _extract_content(self, docx: DocxDocument, parsed_doc: ParsedDocument) -> None:
        """
        提取文档内容（段落和表格）

        遍历文档的所有元素，识别是段落还是表格
        """
        # 遍历文档的所有元素
        element_index = 0
        for element in docx.element.body:
            if isinstance(element, CT_P):
                # 段落
                para = Paragraph(element, docx)
                if para.text.strip():  # 只处理非空段落
                    text_block = self._extract_paragraph(para, element_index)
                    parsed_doc.paragraphs.append(text_block)
                    element_index += 1
            elif isinstance(element, CT_Tbl):
                # 表格
                table = DocxTable(element, docx)
                table_obj = self._extract_table(table, element_index)
                if table_obj.row_count > 0:  # 只处理非空表格
                    parsed_doc.tables.append(table_obj)
                    element_index += 1

    def _extract_paragraph(self, para: Paragraph, index: int) -> TextBlock:
        """
        提取段落及其格式信息

        Args:
            para: docx段落对象
            index: 段落索引

        Returns:
            TextBlock: 文本块对象
        """
        # 提取文本
        text = para.text

        # 提取格式信息
        font_name = None
        font_size = None
        bold = False
        italic = False
        underline = False

        if para.runs:
            # 使用第一个run的格式作为段落格式
            first_run = para.runs[0]
            if first_run.font.name:
                font_name = first_run.font.name
            if first_run.font.size:
                font_size = first_run.font.size.pt
            bold = first_run.bold or False
            italic = first_run.italic or False
            underline = first_run.underline or False

        return TextBlock(
            text=text,
            type="paragraph",
            location=f"第{index + 1}段",
            metadata={
                "index": index,
                "alignment": para.alignment,
                "style": para.style.name if para.style else None,
            },
            font_name=font_name,
            font_size=font_size,
            bold=bold,
            italic=italic,
            underline=underline,
        )

    def _extract_table(self, table: DocxTable, index: int) -> Table:
        """
        提取表格及其内容

        Args:
            table: docx表格对象
            index: 表格索引

        Returns:
            Table: 表格对象
        """
        table_obj = Table(location=f"第{index + 1}个表格")

        # 提取行和单元格
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for col_idx, cell in enumerate(row.cells):
                cell_obj = self._extract_cell(cell, row_idx, col_idx)
                row_cells.append(cell_obj)
            table_obj.rows.append(row_cells)

        # 提取表格元数据
        table_obj.metadata = {
            "row_count": table_obj.row_count,
            "col_count": table_obj.col_count,
            "style": table.style.name if table.style else None,
        }

        return table_obj

    def _extract_cell(self, cell: _Cell, row: int, col: int) -> TableCell:
        """
        提取单元格内容

        Args:
            cell: docx单元格对象
            row: 行号
            col: 列号

        Returns:
            TableCell: 单元格对象
        """
        text = cell.text

        # 提取单元格格式信息
        metadata = {
            "row": row,
            "col": col,
        }

        # 提取格式（使用第一个段落第一个run的格式）
        if cell.paragraphs and cell.paragraphs[0].runs:
            first_run = cell.paragraphs[0].runs[0]
            if first_run.font.name:
                metadata["font_name"] = first_run.font.name
            if first_run.font.size:
                metadata["font_size"] = first_run.font.size.pt
            metadata["bold"] = first_run.bold or False
            metadata["italic"] = first_run.italic or False
            metadata["underline"] = first_run.underline or False

        return TableCell(text=text, row=row, col=col, metadata=metadata)

    def _extract_metadata(self, docx: DocxDocument) -> dict:
        """
        提取文档元数据

        Args:
            docx: docx文档对象

        Returns:
            dict: 元数据字典
        """
        props = docx.core_properties

        metadata = {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "keywords": props.keywords or "",
            "comments": props.comments or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "last_modified_by": props.last_modified_by or "",
            "revision": props.revision or "",
            "category": props.category or "",
        }

        # 添加统计信息
        metadata["paragraph_count"] = len(docx.paragraphs)
        metadata["table_count"] = len(docx.tables)

        return metadata

    def get_text_content(self, parsed_doc: ParsedDocument) -> List[str]:
        """
        获取文档中的所有文本内容

        Args:
            parsed_doc: 解析后的文档对象

        Returns:
            List[str]: 文本列表
        """
        texts = []

        # 添加段落文本
        for para in parsed_doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

        # 添加表格单元格文本
        for table in parsed_doc.tables:
            for row in table.rows:
                for cell in row:
                    if cell.text.strip():
                        texts.append(cell.text.strip())

        return texts

    def get_full_text(self, parsed_doc: ParsedDocument, join_str: str = "\n\n") -> str:
        """
        获取文档的完整文本内容

        Args:
            parsed_doc: 解析后的文档对象
            join_str: 连接字符串

        Returns:
            str: 完整文本
        """
        texts = self.get_text_content(parsed_doc)
        return join_str.join(texts)


def parse_document(file_path: str, extract_comments: bool = False) -> ParsedDocument:
    """
    便捷函数：解析Word文档

    Args:
        file_path: Word文档路径
        extract_comments: 是否提取批注

    Returns:
        ParsedDocument: 解析后的文档对象
    """
    parser = DocumentParser(extract_comments=extract_comments)
    return parser.parse(file_path)
