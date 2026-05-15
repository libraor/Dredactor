"""文档导出器 - 生成脱敏后的Word文档，保留原始格式"""

import os
from typing import Optional, List, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table as DocxTable
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt

from .models import ParsedDocument, TextBlock, Table, TableCell
from .logger import get_logger

logger = get_logger(__name__)


class DocumentExporter:
    """文档导出器

    功能：
    - 生成脱敏后的Word文档
    - 保留原始格式（直接修改原始文档）
    - 支持多种输出模式
    """

    def __init__(self, preserve_formatting: bool = True):
        """
        初始化导出器

        Args:
            preserve_formatting: 是否保留原始格式
        """
        self.preserve_formatting = preserve_formatting

    def export(
        self,
        parsed_doc: ParsedDocument,
        output_path: str,
        overwrite: bool = False,
        original_file_path: Optional[str] = None,
    ) -> bool:
        """
        导出文档到Word文件

        Args:
            parsed_doc: 解析后的文档对象
            output_path: 输出文件路径
            overwrite: 是否覆盖已存在的文件
            original_file_path: 原始文档路径（用于保留格式）

        Returns:
            bool: 是否导出成功
        """
        if os.path.exists(output_path) and not overwrite:
            from .exceptions import FileExistsError_
            raise FileExistsError_(output_path)

        try:
            if self.preserve_formatting and (original_file_path or parsed_doc.file_path):
                source_path = original_file_path or parsed_doc.file_path
                if os.path.exists(source_path):
                    return self._export_with_original_format(parsed_doc, output_path, source_path)

            return self._export_new_document(parsed_doc, output_path)

        except Exception as e:
            from .exceptions import ExportError
            logger.error("导出文档失败: %s", e)
            raise ExportError(output_path, str(e)) from e

    def _export_with_original_format(
        self,
        parsed_doc: ParsedDocument,
        output_path: str,
        original_file_path: str,
    ) -> bool:
        """
        基于原始文档导出，保留所有格式

        Args:
            parsed_doc: 解析后的文档对象（包含修改后的文本）
            output_path: 输出文件路径
            original_file_path: 原始文档路径

        Returns:
            bool: 是否导出成功
        """
        try:
            doc = Document(original_file_path)

            para_idx = 0
            table_idx = 0

            for element in doc.element.body:
                if isinstance(element, CT_P):
                    para = Paragraph(element, doc)
                    if para.text.strip():
                        if para_idx < len(parsed_doc.paragraphs):
                            new_text = parsed_doc.paragraphs[para_idx].text
                            self._update_paragraph_text(para, new_text)
                            para_idx += 1
                elif isinstance(element, CT_Tbl):
                    table = DocxTable(element, doc)
                    if table_idx < len(parsed_doc.tables):
                        self._update_table_content(table, parsed_doc.tables[table_idx])
                        table_idx += 1

            doc.save(output_path)
            return True

        except Exception as e:
            from .exceptions import ExportError
            logger.error("保留格式导出失败: %s", e)
            raise ExportError(output_path, str(e)) from e

    def _update_paragraph_text(self, para: Paragraph, new_text: str) -> None:
        """
        更新段落文本，保留格式

        Args:
            para: docx段落对象
            new_text: 新文本内容
        """
        if not para.runs:
            para.add_run(new_text)
            return

        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        bold = first_run.bold
        italic = first_run.italic
        underline = first_run.underline
        font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None

        for run in para.runs:
            run.text = ""

        first_run.text = new_text

        if font_name:
            first_run.font.name = font_name
        if font_size:
            first_run.font.size = font_size
        if bold is not None:
            first_run.bold = bold
        if italic is not None:
            first_run.italic = italic
        if underline is not None:
            first_run.underline = underline
        if font_color:
            first_run.font.color.rgb = font_color

    def _update_table_content(self, docx_table: DocxTable, parsed_table: Table) -> None:
        """
        更新表格内容，保留格式

        Args:
            docx_table: docx表格对象
            parsed_table: 解析后的表格对象
        """
        for row_idx, row in enumerate(docx_table.rows):
            if row_idx >= len(parsed_table.rows):
                break
            for col_idx, cell in enumerate(row.cells):
                if col_idx >= len(parsed_table.rows[row_idx]):
                    break
                new_text = parsed_table.rows[row_idx][col_idx].text
                self._update_cell_text(cell, new_text)

    def _update_cell_text(self, cell, new_text: str) -> None:
        """
        更新单元格文本，保留格式

        Args:
            cell: docx单元格对象
            new_text: 新文本内容
        """
        if not cell.paragraphs:
            cell.add_paragraph(new_text)
            return

        para = cell.paragraphs[0]
        self._update_paragraph_text(para, new_text)

        for para in cell.paragraphs[1:]:
            for run in para.runs:
                run.text = ""

    def _export_new_document(
        self,
        parsed_doc: ParsedDocument,
        output_path: str,
    ) -> bool:
        """
        创建新文档导出（不保留原始格式）

        Args:
            parsed_doc: 解析后的文档对象
            output_path: 输出文件路径

        Returns:
            bool: 是否导出成功
        """
        try:
            doc = Document()

            if parsed_doc.title:
                doc.core_properties.title = parsed_doc.title

            self._export_paragraphs(doc, parsed_doc)
            self._export_tables(doc, parsed_doc)

            doc.save(output_path)
            return True

        except Exception as e:
            from .exceptions import ExportError
            logger.error("创建新文档失败: %s", e)
            raise ExportError(output_path, str(e)) from e

    def _export_paragraphs(self, doc: Document, parsed_doc: ParsedDocument) -> None:
        """导出段落"""
        for para in parsed_doc.paragraphs:
            self._export_paragraph_text(doc, para)

    def _export_paragraph_text(self, doc: Document, para: TextBlock) -> None:
        """导出单个段落文本"""
        p = doc.add_paragraph(para.text)

        if self.preserve_formatting:
            self._apply_paragraph_format(p, para)

    def _export_tables(self, doc: Document, parsed_doc: ParsedDocument) -> None:
        """导出表格"""
        for table in parsed_doc.tables:
            self._export_table(doc, table)

    def _export_table(self, doc: Document, table: Table) -> None:
        """导出单个表格"""
        docx_table = doc.add_table(rows=table.row_count, cols=table.col_count)

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row):
                docx_table.rows[row_idx].cells[col_idx].text = cell.text

    def _apply_paragraph_format(self, p, para: TextBlock) -> None:
        """应用段落格式"""
        if para.metadata.get("alignment"):
            p.alignment = para.metadata["alignment"]

        if para.metadata.get("style"):
            try:
                p.style = para.metadata["style"]
            except KeyError:
                pass

        if len(p.runs) > 0:
            run = p.runs[0]
            self._apply_run_format(run, para)

    def _apply_run_format(self, run, para: TextBlock) -> None:
        """应用run格式"""
        if para.font_name:
            run.font.name = para.font_name

        if para.font_size:
            run.font.size = Pt(para.font_size)

        if para.bold:
            run.bold = True

        if para.italic:
            run.italic = True

        if para.underline:
            run.underline = True

    def export_with_suffix(
        self,
        parsed_doc: ParsedDocument,
        suffix: str = "_redacted",
        overwrite: bool = False,
    ) -> Optional[str]:
        """
        导出文档到原文件目录，添加后缀

        Args:
            parsed_doc: 解析后的文档对象
            suffix: 文件名后缀
            overwrite: 是否覆盖已存在的文件

        Returns:
            Optional[str]: 输出文件路径，失败返回None
        """
        base_path, ext = os.path.splitext(parsed_doc.file_path)
        output_path = f"{base_path}{suffix}{ext}"

        if self.export(parsed_doc, output_path, overwrite):
            return output_path
        else:
            return None

    def export_batch(
        self,
        parsed_docs: list,
        output_dir: str,
        suffix: str = "_redacted",
        overwrite: bool = False,
    ) -> dict:
        """
        批量导出文档

        Args:
            parsed_docs: 解析后的文档对象列表
            output_dir: 输出目录
            suffix: 文件名后缀
            overwrite: 是否覆盖已存在的文件

        Returns:
            dict: 导出结果统计
        """
        os.makedirs(output_dir, exist_ok=True)

        results = {
            "total": len(parsed_docs),
            "success": 0,
            "failed": 0,
            "files": [],
        }

        for parsed_doc in parsed_docs:
            base_name, ext = os.path.splitext(parsed_doc.file_name)
            output_path = os.path.join(output_dir, f"{base_name}{suffix}{ext}")

            if self.export(parsed_doc, output_path, overwrite):
                results["success"] += 1
                results["files"].append({
                    "input": parsed_doc.file_path,
                    "output": output_path,
                    "success": True,
                })
            else:
                results["failed"] += 1
                results["files"].append({
                    "input": parsed_doc.file_path,
                    "output": output_path,
                    "success": False,
                })

        return results


def export_document(
    parsed_doc: ParsedDocument,
    output_path: str,
    preserve_formatting: bool = True,
    overwrite: bool = False,
    original_file_path: Optional[str] = None,
) -> bool:
    """
    便捷函数：导出文档

    Args:
        parsed_doc: 解析后的文档对象
        output_path: 输出文件路径
        preserve_formatting: 是否保留格式
        overwrite: 是否覆盖已存在的文件
        original_file_path: 原始文档路径（用于保留格式）

    Returns:
        bool: 是否导出成功
    """
    exporter = DocumentExporter(preserve_formatting=preserve_formatting)
    return exporter.export(parsed_doc, output_path, overwrite, original_file_path)
