"""Dredactor CLI主入口"""

import os
import sys
from pathlib import Path
from typing import List, Optional

import typer

from dredactor import (
    DocumentParser,
    RedactionEngine,
    DocumentExporter,
    ReportGenerator,
    RuleManager,
    load_rules,
    export_document,
    generate_report,
    RedactionMapper,
    create_mapper,
    AI_AVAILABLE,
)

if AI_AVAILABLE:
    from dredactor import create_ai_redactor

app = typer.Typer(
    name="dredactor",
    help="Word文档脱敏工具 - 强大的敏感信息识别与替换",
    add_completion=False,
)

@app.command("process", help="处理单个Word文档")
def process_document(
    file_path: str = typer.Argument(..., help="Word文档路径"),
    output: Optional[str] = typer.Option(None, "--output", "-o", help="输出文件路径"),
    rule_group: Optional[str] = typer.Option(None, "--rule-group", "-g", help="使用的规则组"),
    mode: str = typer.Option(
        "mask", "--mode", "-m", help="脱敏模式: replace, mask, partial"
    ),
    replacement: str = typer.Option(
        "****", "--replacement", help="替换文本（replace模式）"
    ),
    override_mode: bool = typer.Option(
        False, "--override-mode", help="是否覆盖规则的默认模式"
    ),
    save_map: bool = typer.Option(
        False, "--save-map", help="保存脱敏映射文件"
    ),
):
    """处理单个Word文档"""
    try:
        # 验证文件存在
        if not os.path.exists(file_path):
            print(f"错误：文件不存在 - {file_path}")
            raise typer.Exit(1)

        # 加载规则
        rule_manager = load_rules()
        engine_rules = []

        if rule_group:
            engine_rules = rule_manager.get_group_rules(rule_group)
            print(f"使用规则组: {rule_group}")
        else:
            engine_rules = rule_manager.get_enabled_rules()
            print(f"使用已启用规则: {len(engine_rules)}条")

        # 创建解析器
        parser = DocumentParser()
        print(f"正在解析文档: {file_path}")
        document = parser.parse(file_path)

        # 创建脱敏引擎
        engine = RedactionEngine(
            rules=engine_rules,
            default_mode=mode,
            replacement_text=replacement,
            override_mode=override_mode,
        )

        # 执行脱敏
        print("正在执行脱敏...")
        result = engine.redact_document(document)

        # 如果需要保存映射，使用特征码方案
        map_data = None
        if save_map:
            mapper = create_mapper()
            map_id = mapper.generate_map_id()

            # 创建映射
            map_data = mapper.create_map(
                map_id=map_id,
                source_file=file_path,
                redacted_file="",  # 稍后设置
                records=result.stats.records,
            )

            # 生成带特征码的文档
            redacted_doc = mapper.apply_markers_to_document(result.redacted_doc, map_data)

            # 生成输出文件名（包含 MAP_ID）
            if not output:
                output = mapper.generate_redacted_filename(file_path, map_id)

            # 更新映射中的脱敏文件路径
            map_data.redacted_file = output

            # 保存映射
            map_path = os.path.join(mapper.MAPPINGS_DIR, f"{map_id}.json")
            if mapper.save_map(map_data, map_path):
                print(f"映射ID: {map_id}")
                print(f"已保存映射文件: {map_path}")
            else:
                print("保存映射文件失败")
                raise typer.Exit(1)

            # 使用带特征码的文档导出
            redacted_to_export = redacted_doc
        else:
            # 不使用特征码，直接使用脱敏结果
            redacted_to_export = result.redacted_doc

            # 确定输出路径
            if not output:
                base, ext = os.path.splitext(file_path)
                output = f"{base}_redacted{ext}"

        # 导出文档
        print(f"正在导出文档: {output}")
        exporter = DocumentExporter()
        if exporter.export(redacted_to_export, output):
            print(f"成功导出到: {output}")
        else:
            print("导出失败")
            raise typer.Exit(1)

        # 显示映射摘要
        if map_data:
            print("\n" + mapper.get_map_summary(map_data))

        # 显示统计信息
        print("\n" + "=" * 60)
        print("脱敏统计")
        print("=" * 60)
        print(f"处理时间: {result.stats.processing_time:.3f}秒")
        print(f"脱敏总数: {result.stats.total_redacted}")

        if result.stats.rules_used:
            print("\n规则使用情况:")
            for rule_name, count in result.stats.rules_used.items():
                rule = rule_manager.get_rule(rule_name)
                desc = rule.description if rule else rule_name
                print(f"  {desc}: {count}次")

        print("\n处理完成!")

    except Exception as e:
        print(f"错误: {str(e)}")
        raise typer.Exit(1)


@app.command("restore", help="恢复脱敏文档")
def restore_document(
    redacted_file: str = typer.Argument(..., help="脱敏后的文档路径（文件名需包含映射ID）"),
    map_file: Optional[str] = typer.Option(None, "--map-file", "-m", help="映射文件路径（可选，默认从文件名推断）"),
    output: Optional[str] = typer.Option(None, "--output", "-o", help="输出文件路径"),
):
    """恢复脱敏文档

    脱敏文件名格式：xxx_redacted_[MAP_ID].docx
    映射文件自动从 .mappings/[MAP_ID].json 加载
    """
    try:
        # 创建映射器
        mapper = create_mapper()

        # 确定映射文件路径
        if not map_file:
            map_file = mapper.get_map_path_from_redacted_file(redacted_file)
            if not map_file:
                print(f"错误：无法从文件名推断映射ID")
                print(f"脱敏文件名格式应为：xxx_redacted_[MAP_ID].docx")
                raise typer.Exit(1)

        # 验证映射文件存在
        if not os.path.exists(map_file):
            print(f"错误：映射文件不存在 - {map_file}")
            raise typer.Exit(1)

        # 验证脱敏文档存在
        if not os.path.exists(redacted_file):
            print(f"错误：脱敏文档不存在 - {redacted_file}")
            raise typer.Exit(1)

        # 确定输出路径
        if not output:
            base, ext = os.path.splitext(redacted_file)
            output = f"{base}_restored{ext}"

        # 加载文件以显示信息
        map_data = mapper.load_map(map_file)
        print("=" * 60)
        print("恢复脱敏文档")
        print("=" * 60)
        print(f"脱敏文档: {redacted_file}")
        print(f"映射文件: {map_file}")
        print(f"输出文件: {output}")
        print(mapper.get_map_summary(map_data))

        # 执行恢复
        print("\n正在恢复...")
        success, restored_doc = mapper.restore_document(
            redacted_file_path=redacted_file, output_path=output, map_path=map_file
        )

        if success:
            print(f"成功恢复到: {output}")
            print("\n恢复完成!")
        else:
            print("恢复失败")
            raise typer.Exit(1)

    except Exception as e:
        print(f"错误: {str(e)}")
        raise typer.Exit(1)


@app.command("demo", help="运行演示")
def run_demo():
    """运行功能演示"""
    print("正在创建演示文档...")

    from docx import Document

    demo_path = "demo_document.docx"
    doc = Document()

    doc.add_heading("Dredactor 演示文档", level=1)
    doc.add_paragraph("本文档包含各种敏感信息，用于演示脱敏功能。")

    doc.add_heading("个人信息", level=2)
    doc.add_paragraph("姓名：张三")
    doc.add_paragraph("手机号：13812345678")
    doc.add_paragraph("邮箱：zhangsan@example.com")
    doc.add_paragraph("身份证号：110101199001011234")
    doc.add_paragraph("银行卡号：6225880112345678")

    doc.add_heading("公司信息", level=2)
    doc.add_paragraph("公司：某某科技有限公司")
    doc.add_paragraph("地址：北京市朝阳区建国路88号")
    doc.add_paragraph("电话：010-12345678")
    doc.add_paragraph("统一社会信用代码：91110000MA00000001")

    doc.add_heading("员工表格", level=2)
    table = doc.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = "姓名"
    table.rows[0].cells[1].text = "电话"
    table.rows[0].cells[2].text = "邮箱"
    table.rows[1].cells[0].text = "李四"
    table.rows[1].cells[1].text = "13987654321"
    table.rows[1].cells[2].text = "lisi@example.com"
    table.rows[2].cells[0].text = "王五"
    table.rows[2].cells[1].text = "13765432198"
    table.rows[2].cells[2].text = "wangwu@example.com"

    doc.save(demo_path)
    print(f"已创建演示文档: {demo_path}")

    # 处理文档
    print("\n正在处理文档...")
    parser = DocumentParser()
    document = parser.parse(demo_path)

    rule_manager = load_rules()
    engine = RedactionEngine(rules=rule_manager.get_enabled_rules())
    result = engine.redact_document(document)

    # 导出
    output_path = "demo_document_redacted.docx"
    exporter = DocumentExporter()
    exporter.export(result.redacted_doc, output_path)

    print(f"已生成脱敏文档: {output_path}")

    # 显示统计
    print("\n脱敏统计:")
    print(f"  总脱敏数: {result.stats.total_redacted}")
    print(f"  处理时间: {result.stats.processing_time:.3f}秒")

    if result.stats.rules_used:
        print("\n规则使用情况:")
        for rule_name, count in result.stats.rules_used.items():
            print(f"    {rule_name}: {count}次")

    print("\n演示完成！可以查看生成的文档:")
    print(f"  - {output_path}")


@app.command("web", help="启动 Web 界面")
def launch_web():
    """启动 Streamlit Web 界面"""


    import subprocess
    import sys

    # 检查 streamlit 是否安装
    try:
        import streamlit
        print("启动 Web 界面...")
        print("=" * 60)
        print("访问地址: http://localhost:8501")
        print("=" * 60)

        # 获取项目根目录
        app_dir = os.path.join(os.path.dirname(__file__), "..", "web")
        app_path = os.path.join(app_dir, "app_simple.py")

        # 运行 streamlit
        subprocess.run(
            [sys.executable, "-m", "streamlit", "run", app_path],
            cwd=os.path.dirname(app_path),
        )
    except ImportError:
        print("错误: Streamlit 未安装")
        print("请运行: pip install streamlit")
        raise typer.Exit(1)


if __name__ == "__main__":
    app()
